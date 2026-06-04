#!/usr/bin/env python3
"""Extract text and tables from a .docx file.

Usage: python3 extract_text.py <docx_file> [output.txt]

Outputs paragraphs as text and tables as TSV blocks.
Requires: pip3 install python-docx
"""
import sys
from docx import Document

def extract_text(docx_path, output_path=None):
    doc = Document(docx_path)
    lines = []
    
    # Extract paragraphs
    para_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
            para_count += 1
    
    lines.append("\n" + "=" * 60)
    lines.append("TABLES")
    lines.append("=" * 60 + "\n")
    
    # Extract tables as TSV
    table_count = 0
    for i, table in enumerate(doc.tables):
        table_count += 1
        lines.append(f"--- Table {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            lines.append('\t'.join(cells))
        lines.append("")
    
    content = '\n'.join(lines)
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"Extracted {para_count} paragraphs + {table_count} tables")
        print(f"Output: {output_path} ({len(content)} chars)")
    else:
        print(content)
    
    return content

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <docx_file> [output.txt]")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None
    extract_text(docx_file, output)
