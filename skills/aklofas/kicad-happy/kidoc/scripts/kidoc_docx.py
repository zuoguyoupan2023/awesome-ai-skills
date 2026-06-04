#!/usr/bin/env python3
"""DOCX generation from kidoc markdown scaffolds.

Converts markdown to DOCX using python-docx.  SVGs are rasterized to
PNG via svglib + rl-renderPM before embedding.  Runs inside reports/.venv/.

Usage (called by kidoc_generate.py, not directly):
    python3 kidoc_docx.py --input reports/HDD.md --output reports/output/HDD.docx
                          --config '{"project": {"name": "..."}}'
"""

from __future__ import annotations

import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

# Add kidoc scripts to path for sibling imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from kidoc_md_parser import parse_markdown
from kidoc_raster import svg_to_png as _svg_to_png, has_svg_render, get_dpi


# ======================================================================
# Inline formatting
# ======================================================================

def _add_runs_to_paragraph(para, runs: list[dict]) -> None:
    """Add formatted runs to a python-docx paragraph."""
    for r in runs:
        run = para.add_run(r['text'])
        if r.get('bold'):
            run.bold = True
        if r.get('italic'):
            run.italic = True
        if r.get('code'):
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xC0, 0x40, 0x00)


# ======================================================================
# Element conversion
# ======================================================================

def _add_element(doc: Document, elem: dict, base_dir: str,
                 dpi: int, temp_files: list) -> None:
    """Add a parsed markdown element to the DOCX document."""
    etype = elem['type']

    if etype == 'heading':
        level = min(elem['level'], 9)
        # python-docx heading levels: 0 = Title, 1-9 = Heading 1-9
        doc_level = 0 if level == 1 else level - 1
        doc.add_heading(elem['text'], level=doc_level)

    elif etype == 'paragraph':
        para = doc.add_paragraph()
        _add_runs_to_paragraph(para, elem['runs'])

    elif etype == 'image':
        _add_image(doc, elem, base_dir, dpi, temp_files)

    elif etype == 'table':
        _add_table(doc, elem)

    elif etype == 'code_block':
        para = doc.add_paragraph(style='No Spacing')
        run = para.add_run(elem['code'])
        run.font.name = 'Courier New'
        run.font.size = Pt(7.5)

    elif etype == 'hr':
        doc.add_paragraph('_' * 50)

    elif etype == 'bullet_list':
        for item_runs in elem['items']:
            para = doc.add_paragraph(style='List Bullet')
            _add_runs_to_paragraph(para, item_runs)

    elif etype == 'numbered_list':
        for item_runs in elem['items']:
            para = doc.add_paragraph(style='List Number')
            _add_runs_to_paragraph(para, item_runs)

    elif etype == 'blockquote':
        para = doc.add_paragraph(style='Quote')
        _add_runs_to_paragraph(para, elem['runs'])


def _add_image(doc: Document, elem: dict, base_dir: str,
               dpi: int, temp_files: list) -> None:
    """Add an image — SVGs are rasterized to PNG first."""
    path = elem['path']
    if not os.path.isabs(path):
        path = os.path.join(base_dir, path)

    if not os.path.isfile(path):
        para = doc.add_paragraph()
        para.add_run(f'[Image not found: {elem["path"]}]').italic = True
        return

    img_path = path
    if path.lower().endswith('.svg'):
        png_path = _svg_to_png(path, dpi=dpi)
        if png_path:
            img_path = png_path
            temp_files.append(png_path)
        else:
            para = doc.add_paragraph()
            para.add_run(f'[SVG rendering unavailable: {elem["path"]}]').italic = True
            return

    try:
        doc.add_picture(img_path, width=Inches(6.0))
        # Center the image
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        para = doc.add_paragraph()
        para.add_run(f'[Failed to embed image: {elem["path"]}]').italic = True

    # Caption
    if elem.get('alt'):
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(elem['alt'])
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _add_table(doc: Document, elem: dict) -> None:
    """Add a table to the DOCX with explicit borders."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    headers = elem['headers']
    rows = elem['rows']
    n_cols = len(headers)

    table = doc.add_table(rows=1 + len(rows), cols=n_cols)

    # Set 'Table Grid' style (provides borders in most templates)
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass  # Style not available in this template

    # Explicit border XML — ensures borders render even without the style
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'C0C0C0')
        borders.append(el)
    tblPr.append(borders)

    # Header row with shading
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
        # Header cell shading
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E8E8F0')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ''
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)


# ======================================================================
# Header / footer
# ======================================================================

def _add_header_footer(doc: Document, config: dict) -> None:
    """Add header and footer to all sections."""
    project = config.get('project', {})
    branding = config.get('reports', {}).get('branding', {})
    header_left = branding.get('header_left', project.get('company', ''))
    classification = config.get('reports', {}).get('classification', '')

    # Resolve placeholders
    for key, val in project.items():
        header_left = header_left.replace(f'{{{key}}}', str(val))

    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        if header_left:
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.text = header_left
            para.style.font.size = Pt(8)
            para.style.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # Footer
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('Page ')
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ======================================================================
# Main generation
# ======================================================================

def generate_docx(markdown_path: str, output_path: str, config: dict) -> str:
    """Convert markdown to DOCX.  Returns the output path."""
    with open(markdown_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    elements = parse_markdown(md_text)
    base_dir = os.path.dirname(os.path.abspath(markdown_path))

    # DPI for SVG rasterization
    dpi = config.get('reports', {}).get('rendering', {}).get('schematic_dpi', 300)

    # Create document (from template if available)
    template = config.get('reports', {}).get('branding', {}).get('cover_template')
    if template and os.path.isfile(template):
        doc = Document(template)
    else:
        doc = Document()

    # Track temp files for cleanup
    temp_files: list[str] = []

    try:
        for elem in elements:
            _add_element(doc, elem, base_dir, dpi, temp_files)

        _add_header_footer(doc, config)

        os.makedirs(os.path.dirname(os.path.abspath(output_path)) or '.', exist_ok=True)
        doc.save(output_path)
    finally:
        # Clean up temp PNG files
        for tf in temp_files:
            try:
                os.unlink(tf)
            except OSError:
                pass

    return output_path


def main():
    parser = argparse.ArgumentParser(description='Generate DOCX from markdown')
    parser.add_argument('--input', '-i', required=True,
                        help='Input markdown file')
    parser.add_argument('--output', '-o', required=True,
                        help='Output DOCX file')
    parser.add_argument('--config', '-c', default='{}',
                        help='JSON config string or path to config file')
    args = parser.parse_args()

    if os.path.isfile(args.config):
        with open(args.config) as f:
            config = json.load(f)
    else:
        config = json.loads(args.config)

    output = generate_docx(args.input, args.output, config)
    print(output, file=sys.stderr)


if __name__ == '__main__':
    main()
