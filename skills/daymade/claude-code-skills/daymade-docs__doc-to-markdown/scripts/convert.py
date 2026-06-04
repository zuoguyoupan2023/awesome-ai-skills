#!/usr/bin/env python3
"""
Multi-tool document to markdown converter with intelligent orchestration.

Supports Quick Mode (fast, single tool) and Heavy Mode (best quality, multi-tool merge).
DOCX files get automatic post-processing to fix pandoc artifacts.

Usage:
    # Quick Mode (default) - fast, single best tool
    uv run --with pymupdf4llm --with markitdown scripts/convert.py document.pdf -o output.md

    # Heavy Mode - multi-tool parallel execution with merge
    uv run --with pymupdf4llm --with markitdown scripts/convert.py document.pdf -o output.md --heavy

    # DOCX deep mode - python-docx direct parsing (experimental)
    uv run --with python-docx scripts/convert.py document.docx -o output.md --docx-deep

    # With image extraction
    uv run --with pymupdf4llm scripts/convert.py document.pdf -o output.md --assets-dir ./images

Dependencies:
    - pymupdf4llm: PDF conversion (LLM-optimized)
    - markitdown: PDF/DOCX/PPTX conversion
    - pandoc: DOCX/PPTX conversion (system install: brew install pandoc)
    - python-docx: DOCX deep parsing (optional, for --docx-deep)
"""

import argparse
import json
import re
import subprocess
import sys
import shutil
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class ConversionResult:
    """Result from a single tool conversion."""
    markdown: str
    tool: str
    images: list[str] = field(default_factory=list)
    success: bool = True
    error: str = ""


# ── Post-processing stats ────────────────────────────────────────────────────

@dataclass
class PostProcessStats:
    """Track what the DOCX post-processor fixed."""
    grid_tables_converted: int = 0
    image_paths_fixed: int = 0
    attributes_removed: int = 0
    code_blocks_fixed: int = 0
    escaped_brackets_fixed: int = 0
    double_brackets_fixed: int = 0

    def any_fixes(self) -> bool:
        return any(
            getattr(self, f) > 0
            for f in self.__dataclass_fields__
        )

    def summary(self) -> str:
        parts = []
        if self.grid_tables_converted:
            parts.append(f"grid tables: {self.grid_tables_converted}")
        if self.image_paths_fixed:
            parts.append(f"image paths: {self.image_paths_fixed}")
        if self.attributes_removed:
            parts.append(f"attributes: {self.attributes_removed}")
        if self.code_blocks_fixed:
            parts.append(f"code blocks: {self.code_blocks_fixed}")
        if self.escaped_brackets_fixed:
            parts.append(f"escaped brackets: {self.escaped_brackets_fixed}")
        if self.double_brackets_fixed:
            parts.append(f"double brackets: {self.double_brackets_fixed}")
        return ", ".join(parts) if parts else "no fixes needed"


# ── DOCX post-processing ─────────────────────────────────────────────────────

# Regex patterns compiled once
_RE_GRID_BORDER = re.compile(r"^\+[:=-][-:=]+(?:\+[:=-][-:=]+)*\+$")
_RE_GRID_ROW = re.compile(r"^\|(.+)\|$")
_RE_NESTED_GRID_BORDER = re.compile(r"^\|\s*\+[:=-][-:=]+\+\s*\|$")
_RE_PANDOC_ATTR = re.compile(r"\{[^}]*(?:width|height)\s*=\s*\"[^\"]*\"[^}]*\}")
_RE_PANDOC_CLASS = re.compile(r"\{\.(?:underline|mark)\}")
_RE_DOUBLE_BRACKET_LINK = re.compile(r"\[\[([^\]]+)\]\(([^)]+)\)")
_RE_DOUBLE_BRACKET_CLOSED = re.compile(r"\[\[([^\]]+)\]\]\(([^)]+)\)")
_RE_DOUBLE_BRACKET_ATTR_LINK = re.compile(r"\[\[([^\]]+)\]\{[^}]*\}\]\(([^)]+)\)")
_RE_ESCAPED_BRACKET = re.compile(r"\\(\[|])")
# Matches single-column dashed line: "  ------"
# AND multi-column simple table border: "  ---- -----"
_RE_DASHED_LINE = re.compile(r"^(\s{2,})-{3,}[\s-]*$")
_RE_ESCAPED_QUOTE = re.compile(r'\\"')
# CJK + fullwidth punctuation range for bold spacing checks
_RE_CJK_PUNCT = re.compile(r'[\u4e00-\u9fff\u3000-\u303f\uff01-\uffef，。、；：！？（）【】「」《》""'']')
_RE_BOLD_PAIR = re.compile(r'\*\*(.+?)\*\*')


def _is_grid_border(line: str) -> bool:
    """Check if a line is a grid table border like +---+ or +:---+."""
    stripped = line.strip()
    return bool(_RE_GRID_BORDER.match(stripped))


def _is_nested_grid_border(line: str) -> bool:
    """Check if a line is a nested grid border like | +---+ |."""
    stripped = line.strip()
    return bool(_RE_NESTED_GRID_BORDER.match(stripped))


def _count_grid_columns(border_line: str) -> int:
    """Count columns in a grid table border line."""
    stripped = border_line.strip()
    if not stripped.startswith("+"):
        return 0
    # Count + separators minus 1 = number of columns
    return stripped.count("+") - 1



# Languages recognized as code block hints in pandoc dashed-line blocks
_KNOWN_CODE_LANGS = frozenset({
    "json", "bash", "shell", "python", "javascript", "js",
    "html", "css", "yaml", "xml", "sql", "plain text",
    "text", "plaintext", "typescript", "ts", "go", "rust",
    "java", "c", "cpp", "ruby", "php",
})


def _build_pipe_table(rows: list[list[str]]) -> list[str]:
    """Build a standard markdown pipe table from rows of cells."""
    if not rows:
        return []
    col_count = max(len(r) for r in rows)
    lines = [
        "| " + " | ".join([""] * col_count) + " |",
        "| " + " | ".join(["---"] * col_count) + " |",
    ]
    for row in rows:
        padded = row + [""] * (col_count - len(row))
        lines.append("| " + " | ".join(padded) + " |")
    return lines


def _collect_images(directory: Path) -> list[str]:
    """Collect image files from a directory (single glob pass)."""
    if not directory.exists():
        return []
    image_exts = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
    return sorted(
        str(p) for p in directory.rglob("*")
        if p.suffix.lower() in image_exts
    )


def _convert_grid_tables(text: str, stats: PostProcessStats) -> str:
    """Convert pandoc grid tables to standard markdown.

    Single-column grid tables (info boxes) -> blockquotes.
    Multi-column grid tables (side-by-side images) -> split into individual elements.
    Nested grid tables are flattened.
    """
    lines = text.split("\n")
    result = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Detect grid table start
        if _is_grid_border(line):
            # Collect the entire grid table
            table_lines = [line]
            i += 1
            while i < len(lines):
                table_lines.append(lines[i])
                if _is_grid_border(lines[i]) and len(table_lines) > 1:
                    i += 1
                    break
                i += 1
            else:
                # Reached end of file without closing border
                # Just output as-is
                result.extend(table_lines)
                continue

            stats.grid_tables_converted += 1
            num_cols = _count_grid_columns(table_lines[0])

            # Extract content lines (skip borders)
            content_lines = []
            for tl in table_lines:
                if _is_grid_border(tl) or _is_nested_grid_border(tl):
                    continue
                m = _RE_GRID_ROW.match(tl.strip())
                if m:
                    content_lines.append(m.group(1).strip())
                else:
                    # Non-standard line inside grid, keep content
                    stripped = tl.strip()
                    if stripped and stripped != "|":
                        content_lines.append(stripped)

            if num_cols <= 1:
                # Single column -> blockquote
                result.append("")
                for cl in content_lines:
                    # Strip outer pipes if present from nested grids
                    cleaned = cl.strip()
                    if cleaned.startswith("|") and cleaned.endswith("|"):
                        cleaned = cleaned[1:-1].strip()
                    # Skip nested grid borders
                    if _RE_GRID_BORDER.match(cleaned):
                        continue
                    if cleaned:
                        result.append(f"> {cleaned}")
                    else:
                        result.append(">")
                result.append("")
            else:
                # Multi-column -> convert to standard pipe table
                # Parse rows: each content_line is a row, split by | into cells
                table_rows = []
                for cl in content_lines:
                    cells = [c.strip() for c in cl.split("|") if c.strip() and not _RE_GRID_BORDER.match(c.strip())]
                    if cells:
                        table_rows.append(cells)

                if table_rows:
                    result.append("")
                    result.extend(_build_pipe_table(table_rows))
                    result.append("")
        else:
            result.append(line)
            i += 1

    return "\n".join(result)


def _fix_image_paths(text: str, assets_dir: Optional[Path], stats: PostProcessStats) -> str:
    """Fix pandoc's double media path and verify images exist.

    Pandoc extracts to <assets_dir>/media/<files> but references as
    <assets_dir>/media/media/<files>. Fix the references.
    Also flatten the actual directory if needed.
    """
    def fix_path(m: re.Match) -> str:
        alt = m.group(1)
        path = m.group(2)
        new_path = path

        # Fix double media/ path
        if "media/media/" in path:
            new_path = path.replace("media/media/", "media/")
            stats.image_paths_fixed += 1

        return f"![{alt}]({new_path})"

    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", fix_path, text)

    # Flatten double media/ nesting if present (pandoc artifact)
    if assets_dir:
        double_media = assets_dir / "media" / "media"
        single_media = assets_dir / "media"
        try:
            for f in double_media.iterdir():
                dest = single_media / f.name
                if not dest.exists():
                    shutil.move(str(f), str(dest))
            double_media.rmdir()
        except (FileNotFoundError, OSError):
                pass

    return text


def _clean_pandoc_attributes(text: str, stats: PostProcessStats) -> str:
    """Remove pandoc attribute annotations from markdown.

    Removes: {width="..." height="..."}, {.underline}, {.mark}, etc.
    """
    count_before = len(text)

    # Remove width/height attributes on images
    text = _RE_PANDOC_ATTR.sub("", text)

    # Remove class attributes like {.underline}
    text = _RE_PANDOC_CLASS.sub("", text)

    if len(text) != count_before:
        # Rough count of removals
        stats.attributes_removed = count_before - len(text)

    return text


def _is_code_content(lines: list[str]) -> bool:
    """Heuristic: decide if content between dashed lines is code or a note/callout.

    Code indicators:
    - Has a language hint on the first line
    - Contains JSON/code-like syntax ({, }, =, ;, ->, //)
    - Contains URLs with protocols
    - Has backslash line continuations

    Note indicators:
    - Mostly CJK/prose text without code syntax
    - Short single-line content
    """
    text = "\n".join(lines)
    stripped = text.strip()

    if not stripped:
        return False

    # Code syntax indicators
    code_chars = set('{}[]();=<>/\\')
    code_char_count = sum(1 for c in stripped if c in code_chars)

    # If >5% of content is code syntax characters, treat as code
    if len(stripped) > 0 and code_char_count / len(stripped) > 0.05:
        return True

    # JSON-like structure
    if stripped.startswith("{") or stripped.startswith("["):
        return True

    # Command-like (starts with common command patterns)
    first_line = lines[0].strip() if lines else ""
    if re.match(r"^(curl|wget|npm|pip|brew|apt|docker|git|ssh|cd|ls|cat|echo|python|node|uv)\s", first_line):
        return True

    return False


def _fix_code_blocks(text: str, stats: PostProcessStats) -> str:
    """Convert pandoc's indented dashed-line blocks to fenced code blocks or blockquotes.

    Pandoc wraps both code and notes in:
      ------------------------------------------------------------------
      content here

      ------------------------------------------------------------------

    With language hint -> code block:
      ```json
      content here
      ```

    Without language hint + prose content -> blockquote:
      > content here

    Without language hint + code-like content -> code block:
      ```
      content here
      ```
    """
    lines = text.split("\n")
    result = []
    i = 0

    known_langs = _KNOWN_CODE_LANGS

    while i < len(lines):
        line = lines[i]

        # Detect indented dashed line (2+ leading spaces, 3+ dashes)
        if _RE_DASHED_LINE.match(line):
            # Check if this is a pandoc simple table (multiple dashed columns
            # on the same line, or content between dashes contains images)
            # Simple table pattern: "  ----  ----" (multiple dash groups separated by spaces)
            # Gap can be 1+ spaces (pandoc uses varying gaps)
            dash_parts = [p for p in line.strip().split() if p.strip()]
            is_simple_table_border = len(dash_parts) > 1 and all(
                re.match(r"^-+$", p.strip()) for p in dash_parts
            )

            if is_simple_table_border:
                # This is a pandoc simple table border - collect rows until
                # next simple table border, convert to pipe table
                table_rows = []
                j = i + 1
                while j < len(lines):
                    next_line = lines[j]
                    # Check for closing simple table border
                    next_parts = [p for p in next_line.strip().split() if p.strip()]
                    is_next_border = len(next_parts) > 1 and all(
                        re.match(r"^-+$", p.strip()) for p in next_parts
                    )
                    if is_next_border:
                        j += 1
                        break
                    if next_line.strip():
                        # Split by 2+ spaces to get columns (pandoc uses varying gaps)
                        cells = [c.strip() for c in re.split(r"\s{2,}", next_line.strip()) if c.strip()]
                        if cells:
                            table_rows.append(cells)
                    j += 1

                if table_rows:
                    stats.code_blocks_fixed += 1
                    result.append("")
                    result.extend(_build_pipe_table(table_rows))
                    result.append("")

                i = j
                continue

            # Not a simple table - look for content and closing dashed line
            block_content = []
            lang_hint = ""
            j = i + 1

            while j < len(lines):
                next_line = lines[j]

                if _RE_DASHED_LINE.match(next_line):
                    # Found closing dashed line
                    j += 1
                    break

                block_content.append(next_line)
                j += 1
            else:
                # No closing dashed line found - not a block, keep as-is
                result.append(line)
                i += 1
                continue

            # If content contains images, treat as simple table (single-column)
            has_images = any("![" in cl for cl in block_content)
            if has_images:
                result.append("")
                for cl in block_content:
                    cl = cl.strip()
                    if cl:
                        result.append(cl)
                result.append("")
                i = j
                continue

            # Check if first line is a language hint (e.g., "  JSON\", "  Plain Text\")
            has_lang_hint = False
            if block_content:
                first = block_content[0].strip()
                first_clean = first.rstrip("\\").strip()
                if first_clean.lower() in known_langs:
                    lang_hint = first_clean.lower()
                    if lang_hint in ("plain text", "text", "plaintext"):
                        lang_hint = ""  # No language tag for plain text
                    has_lang_hint = True
                    block_content = block_content[1:]

            # Clean content: remove leading 2-space indent, fix escaped quotes
            cleaned = []
            for cl in block_content:
                if cl.startswith("  "):
                    cl = cl[2:]
                cl = cl.replace('\\"', '"')
                if cl.endswith("\\"):
                    cl = cl[:-1]
                cleaned.append(cl)

            # Remove trailing/leading empty lines
            while cleaned and not cleaned[-1].strip():
                cleaned.pop()
            while cleaned and not cleaned[0].strip():
                cleaned.pop(0)

            if cleaned:
                stats.code_blocks_fixed += 1

                # Decide: code block vs blockquote
                if has_lang_hint or _is_code_content(cleaned):
                    # Code block — try to pretty-print JSON
                    code_lines = cleaned
                    if lang_hint == "json":
                        try:
                            raw = "\n".join(cleaned)
                            parsed = json.loads(raw)
                            code_lines = json.dumps(parsed, indent=2, ensure_ascii=False).split("\n")
                        except (json.JSONDecodeError, ValueError):
                            pass  # Keep original if not valid JSON

                    result.append("")
                    result.append(f"```{lang_hint}")
                    result.extend(code_lines)
                    result.append("```")
                    result.append("")
                else:
                    # Note/callout -> blockquote
                    result.append("")
                    for cl in cleaned:
                        if cl.strip():
                            result.append(f"> {cl}")
                        else:
                            result.append(">")
                    result.append("")

            i = j
        else:
            result.append(line)
            i += 1

    return "\n".join(result)


def _fix_escaped_brackets(text: str, stats: PostProcessStats) -> str:
    r"""Fix pandoc's escaped brackets: \[ -> [, \] -> ]."""
    count = len(_RE_ESCAPED_BRACKET.findall(text))
    if count:
        stats.escaped_brackets_fixed = count
        text = _RE_ESCAPED_BRACKET.sub(r"\1", text)
    return text


def _fix_double_bracket_links(text: str, stats: PostProcessStats) -> str:
    """Fix double-bracket links: [[text]{.underline}](url) -> [text](url)."""
    count = 0

    def _replace_link(m: re.Match) -> str:
        nonlocal count
        count += 1
        return f"[{m.group(1)}]({m.group(2)})"

    text = _RE_DOUBLE_BRACKET_ATTR_LINK.sub(_replace_link, text)
    text = _RE_DOUBLE_BRACKET_CLOSED.sub(_replace_link, text)
    text = _RE_DOUBLE_BRACKET_LINK.sub(_replace_link, text)

    stats.double_brackets_fixed = count
    return text


def _fix_cjk_bold_spacing(text: str) -> str:
    """Add space around **bold** spans that contain CJK characters.

    DOCX uses run-level styling for bold — no spaces between runs in CJK text.
    Markdown renderers need whitespace around ** to recognize bold boundaries.

    Rule: if a **content** span contains any CJK character, ensure both sides
    have a space (unless already spaced or at line boundary). This handles:
    - CJK directly touching **: 打开**飞书** → 打开 **飞书**
    - Emoji touching **: **密码】**➡️ → **密码】** ➡️
    - Already spaced: 已有 **粗体** → unchanged
    - English bold: English **bold** text → unchanged
    """
    result = []
    last_end = 0

    for m in _RE_BOLD_PAIR.finditer(text):
        start, end = m.start(), m.end()
        content = m.group(1)

        result.append(text[last_end:start])

        # Only add spaces for bold spans containing CJK
        if _RE_CJK_PUNCT.search(content):
            # Space before ** if previous char is not whitespace
            if start > 0 and text[start - 1] not in (' ', '\t', '\n'):
                result.append(' ')

            result.append(m.group(0))

            # Space after ** if next char is not whitespace
            if end < len(text) and text[end] not in (' ', '\t', '\n'):
                result.append(' ')
        else:
            result.append(m.group(0))

        last_end = end

    result.append(text[last_end:])
    return ''.join(result)


def _cleanup_excessive_blank_lines(text: str) -> str:
    """Collapse 3+ consecutive blank lines to 2."""
    return re.sub(r"\n{4,}", "\n\n\n", text)


def postprocess_docx_markdown(
    text: str,
    assets_dir: Optional[Path] = None,
) -> tuple[str, PostProcessStats]:
    """Apply all DOCX-specific post-processing to pandoc markdown output.

    Returns (cleaned_text, stats).
    """
    stats = PostProcessStats()

    # Order matters: grid tables first (they contain images with attributes)
    text = _convert_grid_tables(text, stats)
    text = _fix_image_paths(text, assets_dir, stats)
    text = _clean_pandoc_attributes(text, stats)
    text = _fix_code_blocks(text, stats)
    text = _fix_double_bracket_links(text, stats)
    text = _fix_escaped_brackets(text, stats)
    text = _fix_cjk_bold_spacing(text)
    text = _cleanup_excessive_blank_lines(text)

    return text, stats


# ── DOCX deep parsing (python-docx) ──────────────────────────────────────────

def convert_with_docx_deep(
    file_path: Path, assets_dir: Optional[Path] = None
) -> ConversionResult:
    """Convert DOCX using python-docx direct parsing (experimental).

    More precise than pandoc for:
    - Table structure preservation
    - Comment extraction
    - Image extraction with position info
    """
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        return ConversionResult(
            markdown="",
            tool="docx-deep",
            success=False,
            error="python-docx not installed. Run: pip install python-docx",
        )

    try:
        doc = Document(str(file_path))
        md_parts = []
        images = []
        image_counter = 0

        # Extract images from docx zip
        if assets_dir:
            assets_dir.mkdir(parents=True, exist_ok=True)
            media_dir = assets_dir / "media"
            media_dir.mkdir(exist_ok=True)

            with zipfile.ZipFile(str(file_path), "r") as zf:
                for name in zf.namelist():
                    if name.startswith("word/media/"):
                        img_name = Path(name).name
                        img_dest = media_dir / img_name
                        with zf.open(name) as src, open(img_dest, "wb") as dst:
                            dst.write(src.read())
                        images.append(str(img_dest))

        # Process paragraphs
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if not text:
                md_parts.append("")
                continue

            # Headings
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                md_parts.append(f"{'#' * level} {text}")
                md_parts.append("")
                continue

            # Check for bold-only paragraphs (often sub-headings in Chinese docs)
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if all_bold and para.runs and len(text) < 100:
                md_parts.append(f"**{text}**")
                md_parts.append("")
                continue

            # Regular paragraph
            md_parts.append(text)
            md_parts.append("")

        # Process tables
        for table in doc.tables:
            md_parts.append("")
            rows = table.rows
            if not rows:
                continue

            # Header row
            header_cells = [cell.text.strip() for cell in rows[0].cells]
            md_parts.append("| " + " | ".join(header_cells) + " |")
            md_parts.append("| " + " | ".join(["---"] * len(header_cells)) + " |")

            # Data rows
            for row in rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                md_parts.append("| " + " | ".join(cells) + " |")
            md_parts.append("")

        markdown = "\n".join(md_parts)

        return ConversionResult(
            markdown=markdown,
            tool="docx-deep",
            images=images,
            success=True,
        )
    except Exception as e:
        return ConversionResult(
            markdown="", tool="docx-deep", success=False, error=str(e)
        )


# ── Existing tool converters ─────────────────────────────────────────────────

def check_tool_available(tool: str) -> bool:
    """Check if a conversion tool is available."""
    if tool == "pymupdf4llm":
        try:
            import pymupdf4llm
            return True
        except ImportError:
            return False
    elif tool == "markitdown":
        try:
            import markitdown
            return True
        except ImportError:
            return False
    elif tool == "pandoc":
        return shutil.which("pandoc") is not None
    elif tool == "docx-deep":
        try:
            from docx import Document
            return True
        except ImportError:
            return False
    return False


def select_tools(file_path: Path, mode: str) -> list[str]:
    """Select conversion tools based on file type and mode."""
    ext = file_path.suffix.lower()

    # Tool preferences by format
    tool_map = {
        ".pdf": {
            "quick": ["pymupdf4llm", "markitdown"],  # fallback order
            "heavy": ["pymupdf4llm", "markitdown"],
        },
        ".docx": {
            "quick": ["pandoc", "markitdown"],
            "heavy": ["pandoc", "markitdown"],
        },
        ".doc": {
            "quick": ["pandoc", "markitdown"],
            "heavy": ["pandoc", "markitdown"],
        },
        ".pptx": {
            "quick": ["markitdown", "pandoc"],
            "heavy": ["markitdown", "pandoc"],
        },
        ".xlsx": {
            "quick": ["markitdown"],
            "heavy": ["markitdown"],
        },
    }

    tools = tool_map.get(ext, {"quick": ["markitdown"], "heavy": ["markitdown"]})

    if mode == "quick":
        # Return first available tool
        for tool in tools["quick"]:
            if check_tool_available(tool):
                return [tool]
        return []
    else:  # heavy
        # Return all available tools
        return [t for t in tools["heavy"] if check_tool_available(t)]


def convert_with_pymupdf4llm(
    file_path: Path, assets_dir: Optional[Path] = None
) -> ConversionResult:
    """Convert using PyMuPDF4LLM (best for PDFs)."""
    try:
        import pymupdf4llm

        kwargs = {}
        images = []

        if assets_dir:
            assets_dir.mkdir(parents=True, exist_ok=True)
            kwargs["write_images"] = True
            kwargs["image_path"] = str(assets_dir)
            kwargs["dpi"] = 150

        # Use best table detection strategy
        kwargs["table_strategy"] = "lines_strict"

        md_text = pymupdf4llm.to_markdown(str(file_path), **kwargs)

        if assets_dir:
            images = _collect_images(assets_dir)

        return ConversionResult(
            markdown=md_text, tool="pymupdf4llm", images=images, success=True
        )
    except Exception as e:
        return ConversionResult(
            markdown="", tool="pymupdf4llm", success=False, error=str(e)
        )


def convert_with_markitdown(
    file_path: Path, assets_dir: Optional[Path] = None
) -> ConversionResult:
    """Convert using markitdown."""
    try:
        # markitdown CLI approach
        result = subprocess.run(
            ["markitdown", str(file_path)],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            return ConversionResult(
                markdown="",
                tool="markitdown",
                success=False,
                error=result.stderr,
            )

        return ConversionResult(
            markdown=result.stdout, tool="markitdown", success=True
        )
    except FileNotFoundError:
        # Try Python API
        try:
            from markitdown import MarkItDown

            md = MarkItDown()
            result = md.convert(str(file_path))
            return ConversionResult(
                markdown=result.text_content, tool="markitdown", success=True
            )
        except Exception as e:
            return ConversionResult(
                markdown="", tool="markitdown", success=False, error=str(e)
            )
    except Exception as e:
        return ConversionResult(
            markdown="", tool="markitdown", success=False, error=str(e)
        )


def convert_with_pandoc(
    file_path: Path, assets_dir: Optional[Path] = None
) -> ConversionResult:
    """Convert using pandoc.

    Pandoc's --extract-media=DIR creates a media/ subdirectory inside DIR.
    We point --extract-media at assets_dir's parent so pandoc's media/
    subdirectory lands exactly at assets_dir (when assets_dir ends with 'media'),
    or we use a temp dir and move files afterward.
    """
    try:
        cmd = ["pandoc", str(file_path), "-t", "markdown", "--wrap=none"]

        extract_dir = None
        if assets_dir:
            assets_dir.mkdir(parents=True, exist_ok=True)
            # Pandoc always creates a media/ subdirectory inside --extract-media.
            # Point it at the parent so media/ lands at assets_dir.
            if assets_dir.name == "media":
                extract_dir = assets_dir.parent
            else:
                extract_dir = assets_dir
            cmd.extend(["--extract-media", str(extract_dir)])

        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=120
        )

        if result.returncode != 0:
            return ConversionResult(
                markdown="", tool="pandoc", success=False, error=result.stderr
            )

        md = result.stdout

        # Convert absolute image paths to relative paths based on output location
        if extract_dir:
            abs_media = str(extract_dir / "media")
            # Replace absolute paths with relative 'media/' prefix
            md = md.replace(abs_media + "/", "media/")

        images = _collect_images(assets_dir) if assets_dir else []

        return ConversionResult(
            markdown=md, tool="pandoc", images=images, success=True
        )
    except Exception as e:
        return ConversionResult(
            markdown="", tool="pandoc", success=False, error=str(e)
        )


def convert_single(
    file_path: Path, tool: str, assets_dir: Optional[Path] = None
) -> ConversionResult:
    """Run a single conversion tool."""
    converters = {
        "pymupdf4llm": convert_with_pymupdf4llm,
        "markitdown": convert_with_markitdown,
        "pandoc": convert_with_pandoc,
        "docx-deep": convert_with_docx_deep,
    }

    converter = converters.get(tool)
    if not converter:
        return ConversionResult(
            markdown="", tool=tool, success=False, error=f"Unknown tool: {tool}"
        )

    return converter(file_path, assets_dir)


def merge_results(results: list[ConversionResult]) -> ConversionResult:
    """Merge results from multiple tools, selecting best segments."""
    if not results:
        return ConversionResult(markdown="", tool="none", success=False)

    # Filter successful results
    successful = [r for r in results if r.success and r.markdown.strip()]
    if not successful:
        # Return first error
        return results[0] if results else ConversionResult(
            markdown="", tool="none", success=False
        )

    if len(successful) == 1:
        return successful[0]

    # Multiple successful results - merge them
    # Strategy: Compare key metrics and select best
    best = successful[0]
    best_score = score_markdown(best.markdown)

    for result in successful[1:]:
        score = score_markdown(result.markdown)
        if score > best_score:
            best = result
            best_score = score

    # Merge images from all results
    all_images = []
    seen = set()
    for result in successful:
        for img in result.images:
            if img not in seen:
                all_images.append(img)
                seen.add(img)

    best.images = all_images
    best.tool = f"merged({','.join(r.tool for r in successful)})"

    return best


def score_markdown(md: str) -> float:
    """Score markdown quality for comparison."""
    score = 0.0

    # Length (more content is generally better)
    score += min(len(md) / 10000, 5.0)  # Cap at 5 points

    # Tables (proper markdown tables)
    table_count = md.count("|---|") + md.count("| ---")
    score += min(table_count * 0.5, 3.0)

    # Images (referenced images)
    image_count = md.count("![")
    score += min(image_count * 0.3, 2.0)

    # Headings (proper hierarchy)
    h1_count = md.count("\n# ")
    h2_count = md.count("\n## ")
    h3_count = md.count("\n### ")
    if h1_count > 0 and h2_count >= h1_count:
        score += 1.0  # Good hierarchy

    # Lists (structured content)
    list_count = md.count("\n- ") + md.count("\n* ") + md.count("\n1. ")
    score += min(list_count * 0.1, 2.0)

    # Penalize pandoc artifacts (grid tables, attributes)
    artifact_count = md.count("+:---") + md.count("+---+")
    artifact_count += md.count('{width="') + md.count("{.underline}")
    score -= artifact_count * 0.5

    return score


def main():
    parser = argparse.ArgumentParser(
        description="Convert documents to markdown with multi-tool orchestration",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
    # Quick mode (default)
    python convert.py document.pdf -o output.md

    # Heavy mode (best quality)
    python convert.py document.pdf -o output.md --heavy

    # DOCX deep mode (python-docx parsing)
    python convert.py document.docx -o output.md --docx-deep

    # With custom assets directory
    python convert.py document.pdf -o output.md --assets-dir ./images
        """,
    )
    parser.add_argument("input", type=Path, nargs="?", help="Input document path")
    parser.add_argument(
        "-o", "--output", type=Path, help="Output markdown file"
    )
    parser.add_argument(
        "--heavy",
        action="store_true",
        help="Enable Heavy Mode (multi-tool, best quality)",
    )
    parser.add_argument(
        "--docx-deep",
        action="store_true",
        help="Use python-docx direct parsing (experimental, DOCX only)",
    )
    parser.add_argument(
        "--no-postprocess",
        action="store_true",
        help="Disable DOCX post-processing (keep raw pandoc output)",
    )
    parser.add_argument(
        "--assets-dir",
        type=Path,
        default=None,
        help="Directory for extracted images (default: <output>_assets/)",
    )
    parser.add_argument(
        "--tool",
        choices=["pymupdf4llm", "markitdown", "pandoc", "docx-deep"],
        help="Force specific tool (overrides auto-selection)",
    )
    parser.add_argument(
        "--list-tools",
        action="store_true",
        help="List available tools and exit",
    )

    args = parser.parse_args()

    # List tools mode
    if args.list_tools:
        tools = ["pymupdf4llm", "markitdown", "pandoc", "docx-deep"]
        print("Available conversion tools:")
        for tool in tools:
            status = "+" if check_tool_available(tool) else "-"
            print(f"  {status} {tool}")
        sys.exit(0)

    # Validate input
    if args.input is None:
        parser.error("the following arguments are required: input")
    if not args.input.exists():
        print(f"Error: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    # Determine output path
    output_path = args.output or args.input.with_suffix(".md")

    # Determine assets directory
    assets_dir = args.assets_dir
    if assets_dir is None:
        assets_dir = output_path.parent / f"{output_path.stem}_assets"

    is_docx = args.input.suffix.lower() in (".docx", ".doc")

    # Handle --docx-deep mode
    if args.docx_deep:
        if not is_docx:
            print("Error: --docx-deep only works with DOCX files.", file=sys.stderr)
            sys.exit(1)
        tools = ["docx-deep"]
    elif args.tool:
        tools = [args.tool] if check_tool_available(args.tool) else []
    else:
        # Select tools
        mode = "heavy" if args.heavy else "quick"
        tools = select_tools(args.input, mode)

    mode = "docx-deep" if args.docx_deep else ("heavy" if args.heavy else "quick")

    if not tools:
        print("Error: No conversion tools available.", file=sys.stderr)
        print("Install with:", file=sys.stderr)
        print("  pip install pymupdf4llm", file=sys.stderr)
        print("  uv tool install markitdown[pdf]", file=sys.stderr)
        print("  brew install pandoc", file=sys.stderr)
        sys.exit(1)

    print(f"Converting: {args.input}")
    print(f"Mode: {mode.upper()}")
    print(f"Tools: {', '.join(tools)}")

    # Run conversions
    results = []
    for tool in tools:
        print(f"  Running {tool}...", end=" ", flush=True)

        # Use separate assets dirs for each tool in heavy mode
        tool_assets = None
        if assets_dir and mode == "heavy" and len(tools) > 1:
            tool_assets = assets_dir / tool
        elif assets_dir:
            tool_assets = assets_dir

        result = convert_single(args.input, tool, tool_assets)
        results.append(result)

        if result.success:
            print(f"ok ({len(result.markdown):,} chars, {len(result.images)} images)")
        else:
            print(f"FAIL ({result.error[:50]}...)")

    # Merge results if heavy mode
    if mode == "heavy" and len(results) > 1:
        print("  Merging results...", end=" ", flush=True)
        final = merge_results(results)
        print(f"ok (using {final.tool})")
    else:
        final = merge_results(results)

    if not final.success:
        print(f"Error: Conversion failed: {final.error}", file=sys.stderr)
        sys.exit(1)

    # Apply DOCX post-processing
    if is_docx and not args.no_postprocess and "pandoc" in final.tool:
        print("  Post-processing DOCX output...", end=" ", flush=True)
        final.markdown, pp_stats = postprocess_docx_markdown(
            final.markdown, assets_dir
        )
        print(f"ok ({pp_stats.summary()})")

    # Write output
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(final.markdown)

    print(f"\nOutput: {output_path}")
    print(f"  Size: {len(final.markdown):,} characters")
    if final.images:
        print(f"  Images: {len(final.images)} extracted")


if __name__ == "__main__":
    main()
